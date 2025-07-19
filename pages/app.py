import streamlit as st
from firebase_config import (
    check_and_increment_usage,
    save_script,
    upgrade_to_pro,
    is_user_pro,
    db
)
from email_utils import generate_otp, send_otp_email
import requests
from fpdf import FPDF
from docx import Document
import io
from firebase_config import save_feedback
# CONFIG
st.set_page_config(page_title="🎬 Scriptify AI", layout="centered")

API_URL = "https://api.together.xyz/v1/completions"
headers = {
    "Authorization": f"Bearer {st.secrets['TOGETHER_API_KEY']}",
    "Content-Type": "application/json"
}

def query_together(prompt):
    payload = {
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "prompt": prompt,
        "max_tokens": 1200,
        "temperature": 0.7,
        "top_p": 0.9
    }
    try:
        response = requests.post(API_URL, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        return response.json().get("choices", [{}])[0].get("text", "⚠️ No response from model.")
    except Exception as e:
        return f"⚠️ Error contacting Together API: {e}"

# ------------- LOGIN -------------
st.title("🎬 Scriptify AI")

if "user_email" not in st.session_state:
    user_email = st.text_input("Enter your email")
    if user_email and not st.session_state.get("otp_sent"):
        otp = generate_otp()
        send_otp_email(user_email, otp)
        st.session_state["otp"] = otp
        st.session_state["otp_sent"] = True
        st.session_state["pending_email"] = user_email
        st.success(f"OTP sent to {user_email}")

    if st.session_state.get("otp_sent"):
        entered_otp = st.text_input("Enter OTP")
        if st.button("Verify OTP"):
            if entered_otp == st.session_state["otp"]:
                st.session_state["user_email"] = st.session_state["pending_email"]
                st.session_state.pop("otp", None)
                st.session_state.pop("otp_sent", None)
                st.session_state.pop("pending_email", None)

                if is_user_pro(st.session_state["user_email"]):
                    st.session_state["is_pro"] = True
                    st.success("OTP Verified. Welcome, Pro user!")
                else:
                    st.session_state["is_pro"] = False
                    st.success("OTP Verified. Welcome!")
            else:
                st.error("Invalid OTP.")
                st.stop()
    if "user_email" not in st.session_state:
        st.stop()

# ------------- SIDEBAR -------------
st.sidebar.success(f"Logged in as: {st.session_state['user_email']}")

if st.sidebar.button("Logout"):
    st.session_state.clear()
    st.success("Logged out. Refresh to login again.")
    st.stop()

if st.session_state.get("is_pro"):
    st.sidebar.markdown("💎 **Pro Plan Activated**")
else:
    if st.sidebar.button("💎 Upgrade to Pro"):
        upgrade_to_pro(st.session_state["user_email"])
        st.session_state["is_pro"] = True
        st.success("Pro activated! You now have access to premium features.")

user_doc = db.collection("users").document(st.session_state["user_email"]).get()
if user_doc.exists:
    data = user_doc.to_dict()
    expiry = data.get("pro_expiry")
    if expiry:
        st.sidebar.markdown(f"🗓️ **Pro valid until:** `{expiry}`")

# ------------- Usage Limit (Free Users) -------------
if not st.session_state.get("is_pro"):
    _, remaining = check_and_increment_usage(st.session_state["user_email"], increment=False)
    st.info(f"Remaining scripts today: **{remaining}/20**")

# ------------- INPUT FORM -------------
st.header("Generate Your Script")
topic = st.text_input("Enter your Video Topic", placeholder="e.g., AI tools for students")
platform = st.selectbox("Select your Platform", ["YouTube", "Instagram Reels"])

tone_options = ["Motivational", "Funny", "Professional"]
if st.session_state.get("is_pro"):
    tone_options += ["Casual", "Inspirational", "Bold"]
tone = st.selectbox("Selete the Tone of the script", tone_options)

audience = st.text_input("Enter your Target Audience", placeholder="e.g., creators, students")

if "script" not in st.session_state:
    st.session_state["script"] = None

# ------------- GENERATE SCRIPT -------------
if st.button("Generate Script"):
    if not st.session_state["is_pro"]:
        allowed, remaining = check_and_increment_usage(st.session_state["user_email"])
        if not allowed:
            st.warning("🚫 Daily limit reached (20). Upgrade to Pro for unlimited access.")
            st.stop()

    with st.spinner("Generating your script..."):
        prompt = f"""
Write a complete {platform} video script for the topic "{topic}" in a {tone} tone. Target audience: {audience or 'general audience'}.

🎬 Opening Scene
📌 Section 1: Introduction
📌 Section 2: Tips or insights
📌 Section 3: Motivation or CTA
🎬 Closing Scene

Make it engaging, clear, and appropriate for the platform.
"""
        script = query_together(prompt)
        st.session_state["script"] = script
        save_script(st.session_state["user_email"], topic, platform, tone, audience, script)
        st.success("Your Script is saved!")

# ------------- DISPLAY & EXPORT -------------
if st.session_state["script"]:
    st.subheader("Your Script")
    st.write(st.session_state["script"])

    if st.session_state.get("is_pro"):
        # PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in st.session_state["script"].split("\n"):
            pdf.multi_cell(0, 10, line.encode("latin-1", "ignore").decode("latin-1"))
        pdf_bytes = io.BytesIO(pdf.output(dest='S').encode('latin-1'))
        pdf_bytes.seek(0)
        st.download_button("Download PDF", data=pdf_bytes, file_name="script.pdf", mime="application/pdf")

        # DOCX
        doc = Document()
        doc.add_heading("Generated Script", 0)
        for line in st.session_state["script"].split("\n"):
            doc.add_paragraph(line)
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        st.download_button("Download DOCX", data=docx_bytes, file_name="script.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # TXT
        txt_bytes = io.BytesIO(st.session_state["script"].encode('utf-8'))
        st.download_button("Download TXT", data=txt_bytes, file_name="script.txt", mime="text/plain")

        # Refinement
        st.subheader("AI Script Refinement")
        refine_prompt = st.text_area("Suggest edits or improvements", placeholder="Make it more engaging...")

        if st.button("Refine Script"):
            with st.spinner("Refining..."):
                refined = query_together(f"Improve this script:\n{st.session_state['script']}\n\nSuggestions: {refine_prompt}")
                st.session_state["script"] = refined
                st.success("Refined Script")
                st.write(refined)

                # Save refined version to Firestore
                save_script(st.session_state["user_email"], topic + " (Refined)", platform, tone, audience, refined)

# ------------- SUPPORT -------------
st.sidebar.markdown("💬 Have Feedback?")
with st.sidebar.form("feedback_form"):
    feedback = st.text_area("Tell us what you think:")
    submit = st.form_submit_button("📩 Submit")
    if submit and feedback.strip():
        save_feedback(st.session_state.get("user_email", "Anonymous"), feedback.strip())
        st.success("✅ Thanks for your feedback!")
    elif submit:
        st.warning("⚠️ Please write something before submitting.")


