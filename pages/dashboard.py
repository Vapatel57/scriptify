from firebase_config import get_user_scripts, upgrade_to_pro, is_user_pro, db
from email_utils import generate_otp, send_otp_email
import streamlit as st
from fpdf import FPDF
from docx import Document
import io

st.set_page_config(page_title="📜 Scriptify AI | Dashboard", layout="centered")
st.title("Your Scripts Dashboard")

# ------------------ LOGIN ------------------
if "user_email" not in st.session_state:
    if "otp_verified" not in st.session_state:
        st.session_state["otp_verified"] = False

    user_email = st.text_input("Enter your email")

    if user_email and not st.session_state.get("otp_sent"):
        otp = generate_otp()
        send_otp_email(user_email, otp)
        st.session_state["otp"] = otp
        st.session_state["otp_sent"] = True
        st.session_state["pending_email"] = user_email
        st.success(f"OTP sent to {user_email}")

    if st.session_state.get("otp_sent") and not st.session_state["otp_verified"]:
        entered_otp = st.text_input("Enter the OTP")
        if st.button("Verify OTP"):
            if entered_otp == st.session_state.get("otp"):
                st.success("OTP Verified. Displaying your scripts.")
                user_email = st.session_state["pending_email"]
                st.session_state["user_email"] = user_email
                st.session_state["otp_verified"] = True
                st.session_state.pop("otp", None)
                st.session_state.pop("otp_sent", None)
                st.session_state.pop("pending_email", None)
                st.session_state["pro_user"] = is_user_pro(user_email)
            else:
                st.error("Incorrect OTP.")
    if "user_email" not in st.session_state:
        st.stop()

# ------------------ LOGOUT ------------------
st.sidebar.write(f"Logged in as: {st.session_state['user_email']}")
if st.sidebar.button("Logout"):
    st.session_state.clear()
    st.success("Logged out. Refresh to log in again.")
    st.stop()

# ------------------ PRO ------------------
if is_user_pro(st.session_state["user_email"]):
    st.session_state["pro_user"] = True
    st.markdown("**Pro Plan Activated** — Enjoy unlimited access!")
else:
    if st.button("✨ Upgrade to Pro"):
        upgrade_to_pro(st.session_state["user_email"])
        st.session_state["pro_user"] = True
        st.success("You're now a Pro user with unlimited access!")

# ------------------ SEARCH BAR ------------------
search = st.text_input("Search by topic or date (yyyy-mm-dd)", "")

# ------------------ SCRIPTS ------------------
scripts = get_user_scripts(st.session_state["user_email"])

if not scripts:
    st.warning("No scripts found.")
else:
    st.subheader("Your Saved Scripts")

    for script in scripts:
        if search.lower() not in str(script['topic']).lower() and search not in str(script['timestamp']):
            continue

        st.markdown(f"---\n### {script['topic']} — `{script['timestamp']}`")
        if script.get("refined"):
            st.markdown("**Refined Script**")

        st.code(script['script'], language="markdown")

        # --- Download buttons ---
        col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
        with col1:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in script['script'].split("\n"):
                pdf.multi_cell(0, 10, line.encode("latin-1", "ignore").decode("latin-1"))
            pdf_bytes = io.BytesIO(pdf.output(dest='S').encode('latin-1'))
            pdf_bytes.seek(0)
            st.download_button("PDF", data=pdf_bytes, file_name="script.pdf", mime="application/pdf", key=f"pdf-{script['timestamp']}")

        with col2:
            doc = Document()
            doc.add_heading(script['topic'], 0)
            for line in script['script'].split("\n"):
                doc.add_paragraph(line)
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            st.download_button("DOCX", data=docx_bytes, file_name="script.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"docx-{script['timestamp']}")

        with col3:
            txt_bytes = io.BytesIO(script['script'].encode("utf-8"))
            st.download_button("TXT", data=txt_bytes, file_name="script.txt", mime="text/plain", key=f"txt-{script['timestamp']}")

        # --- Delete Button ---
        with col4:
            if st.button("Delete", key=f"delete-{script['timestamp']}"):
                docs = (
                    db.collection("scripts")
                    .where("user_email", "==", st.session_state["user_email"])
                    .where("timestamp", "==", script["timestamp"])
                    .stream()
                )
                for doc in docs:
                    doc.reference.delete()
                st.success("Script deleted. Refresh to see changes.")
                st.experimental_rerun()
